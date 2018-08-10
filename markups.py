#!/usr/bin/env python
import docx  ##for microsoft word
import sys
import getopt
from openpyxl import Workbook  ##for microsoft excel
from strippers import strip_address ##strips institution and address
from strippers import strip_email ##strips {} email addresses
from strippers import strip_language ##strips () language
from strippers import strip_original ##strips original language

def main(argv):

    Region='ME'
    Month='July'
    Year='18'

    try:
        opts, args = getopt.getopt(argv,"hi:o:",["Month=","Region=","Year="])
    except getopt.GetoptError:
        usage()
        ##print 'test.py -i <inputfile> -o <outputfile>'
        sys.exit(2)
    for opt, arg in opts:
        if opt == '-h':
            print 'markups.py --Month <Month> --Year <Year> --Region <Region>'
            sys.exit()
        elif opt in ("-m", "--Month"):
            Month = arg
        elif opt in ("-r", "--Region"):            
            Region = arg
        elif opt in ("-Y", "--Year"):            
            Year = arg    

    print 'converting %s abstracts to markups %s 20%s' % (Region, Month, Year)


    ###Opening File

    
    
    Months_date={'January':'31 Jan', 'February':'28 Feb', 'March': '31 March', 'April':'30 April', 'May':'31 May', 'June': '30 June', 'July':'31 July', 'August':'31 August', 'September':'30 Sept', 'October':'31 Oct', 'November':'30 Nov', 'December':'30 Dec'}
    Months_no_dict={'January':1, 'February':2, 'March':3, 'April':4, 'May':5, 'June':6, 'July':7, 'August':8, 'September':9, 'October':10, 'November':11, 'December':12}
    month_no=Months_no_dict[Month]
    Myfilepath='/Users/itadmin/Documents/IAEA/Abstracts/'
    ###Myfilepath='/cygdrive/c/Users/K1343788/Documents/IAEA/Monthly/20%s/' % Year
    ###Myfilepath='/cygdrive/n/My Documents/Documents/Documents/20%s_Abstracts/' % Year 
    abstracts_path= '%sKCL OS updates %s - %02d %s.docx' % (Myfilepath, Region.upper(), Months_no_dict[Month], Year)
    ###abstracts_path='%sKCL.docx' % Myfilepath

    ###print 'trying to find ', abstracts_path
    try:
        doc = docx.Document(abstracts_path)
        print "Opening ", abstracts_path
        length =  len(doc.paragraphs)
    except IOError:
        print 'Cannot open: ', abstract_path, '\n Check filepath and file exists'

    ###Define variables

    Country=''
    abstract_number=0
    Markups={} ##dict for all markups
    markup={} ##dummy dict for each markup
    abstract_reference=''
    Language=''
    authors_dict={} ##dummy dict for authors
    authorsinc=0  ##to include authors over multiple lines
    
    ###Loop through paragraphs in document
    for para in doc.paragraphs:
        #print para.text
        if 'ROK'==para.text.strip():
            Country_index='KR'
            Language='KOREAN'
            Country='ROK'
            abstract_number=0
        elif 'Japan'==para.text.strip():
            Country_index='JP'
            Language='JAPANESE'
            Country='Japan'
            abstract_number=0
        elif 'Indonesia'==para.text.strip():
            Country_index='ID'
            Language='INDONESIAN'
            Country='Indonesia'
            abstract_number=0
            print 'Country= Indonesian'
        elif 'Egypt'==para.text.strip():
            Country_index='EG'
            Language='ARABIC'
            Country='Egypt'
            abstract_number=0
        elif 'Iran'==para.text.strip():
            Country_index='IR'
            Language='FARSI'
            Country='Iran'
            abstract_number=0
        elif 'Turkey'==para.text.strip():
            Country_index='TR'
            Language='TURKISH'
            Country='Turkey'
            abstract_number=0
        elif 'Jordan'==para.text.strip():
            Country_index='JO'
            Language='ARABIC'
            Country='Jordan'
            abstract_number=0
        elif 'Morocco'==para.text.strip():
            Country_index='MA'
            Language='ARABIC'
            Country='Morocco'
            abstract_number=0
        elif 'Tunisia'==para.text.strip():
            Country_index='TN'
            Language='ARABIC'
            Country='Tunisia'
            abstract_number=0
        elif 'Algeria'==para.text.strip():
            Country_index='DZ'
            Language='ARABIC'
            Country='Algeria'
            abstract_number=0
        elif 'Saudi Arabia'==para.text.strip():
            Country_index='SA'
            Language='ARABIC'
            Country='Saudi Arabia'
            abstract_number=0
        elif 'UAE'==para.text.strip():
            Country_index='AE'
            Language='ARABIC'
            Country='UAE'
            abstract_number=0
        elif 'Taiwan'==para.text.strip():
            Country_index='TW'
            Language='MANDARIN'
            Country='Taiwan'
            abstract_number=0
        elif 'Myanmar'==para.text.strip():
            Country_index='MM'
            Language='BURMESE'
            Country='Myanmar'
            abstract_number=0
        elif 'Morocco'==para.text.strip():
            Country_index='MA'
            Language='FRENCH'
            Country='Morocco'
            abstract_number=0    


        elif 'Title:' in para.text:
            ###check for Country
            try:
                Country
            except:
                print 'Country Section Heading Missing'
                exit(2)
            ###Get new abstract number
            abstract_number=abstract_number+1
            authorsinc=0
            abstract_reference='kcl.%s.%02d.20%s.%s.%02d' % (Region.lower(), month_no, Year, Country_index, abstract_number)
            #####Markups[abstract_reference]={}

            ###Empty markup and authors dict
            markup.clear()
            authors_dict.clear()
            Title=para.text
            Title=Title.replace('Title:',"",1)
            ###print Language, Title
            if Language in Title:
                print 'Language=',Language
                Title=Title.replace('(%s)' % Language,"",1)
                Title=Title.strip()
            else:
                print '*** detecting language',
                #print 'spliting', Title.split('(').pop()
                dummy_Language='('+Title.split('(').pop()
                Language=strip_language(dummy_Language)
                print ' -- Language=', Language
                Title=Title.replace('(%s)' % Language,"",1)
                Title=Title.strip()
            if '<<' in Title:
                print 'stripping original language'
                Original_Title=strip_original(Title)
                Title=Title.replace('<<%s>>' % Original_Title,"",1)
                Title=Title.strip()
                print 'Original Title=',Original_Title
            ###print Language, Title    
            markup['Title']=Title
            markup['Language']=Language
            markup['Country']=Country
            print abstract_reference, Title
            #### Set up markup dict
            Markups[abstract_reference]={}
            Markups[abstract_reference]['Title']=markup['Title']  ##Put into markup
            Markups[abstract_reference]['Language']=markup['Language']  ##Put into markup
            Markups[abstract_reference]['Country']=markup['Country']  ##Put into markup

        elif 'Date:' in para.text:
            authorsinc=0
            Date=para.text
            Date=Date.replace('Source Date:',"",1)
            Date=Date.replace('Date:',"",1)
            Date=Date.replace('Accessed',"",1)
            Date=Date.strip()
            markup['Date']=Date
            Markups[abstract_reference]['Date']=markup['Date']  ##Put into markup

        elif 'Source:' in para.text or 'source:' in para.text:
            authorsinc=0
            #print para.text
            Source=para.text
            Source=Source.replace('Original Source:',"",1)
            Source=Source.replace('Original source:',"",1)
            Source=Source.replace('Source:',"",1)
            Source=Source.strip()
            Source=Source.replace('(%s)' % Language, "",1)
            markup['Source']=Source
            ###print 'Source is', Source
            Markups[abstract_reference]['Source']=markup['Source']  ##Put into markup
            
        elif 'Authors:' in para.text or 'Author:' in para.text or authorsinc==1:
            ###control for authors over multiple lines
            if authorsinc==1 and not para.text:
                authorsinc=0
        
            else:
                if authorsinc==0:
                    authorsinc=1
                    
                ###Define dummy variables for author dicts
                authors_dummy_list=[]
                multi_addr_list=[]
                address=''
                email=''
                org=''
                suborg=''
                subsuborg=''
                
                ###Strip intial text
                authors_dummy=para.text
                authors_dummy=authors_dummy.replace('Authors:',"",1)
                authors_dummy=authors_dummy.replace('Author:',"",1)
                #print 'authors_dummy', authors_dummy
                ###Look for author addresses and loop over authors
                while '(' in authors_dummy:
                    address= strip_address(authors_dummy)
                    authors_dummy=authors_dummy.replace('(%s)' % address,"",1)
                    authors_dummy_list=authors_dummy.split(',')
                    #print 'address=',address, 'authors=',authors_dummy
                    ###Loop over authors by address
                    for author_dummy in authors_dummy_list:
                        email=''
                        author=author_dummy
                        
                        ###Find emails
                        if '{' in author_dummy:
                            #print 'finding email', author_dummy
                            email=strip_email(author_dummy)
                            author=author.replace('{%s}' % email,"",1)
                            email=email.strip('email:')
                            email=email.strip('{')
                            email=email.strip('}')
                            email=email.strip()
                            ###end of emails    

                        ###Multiple addresses
                        if '/' in address:
                            ###print 'multi addr', author
                            multi_addr_list=address.split('/')
                            ###print multi_addr_list
                        ###Asign orgs, suborgs, subsuborgs
                        author=author.strip()
                        orgs=address.split(',')
                        org=orgs[0]
                        org=org.strip()
                        
                        if len(orgs)==3:
                            org=orgs[2]
                            suborg=orgs[1]
                            subsuborg=orgs[0]
                        elif len(orgs)==2:
                            org=orgs[1]
                            suborg=orgs[0]
                            subsuborg=''
                        elif len(orgs)==1:
                            org=orgs[0]
                            suborg=''
                            subsuborg=''
                        else:
                            org=''
                            suborg=''
                            subsuborg=''
                            
                        org=org.strip()
                        suborg=suborg.strip()
                        subsuborg=subsuborg.strip()
                        ###print '\nAuthor:', author, 'Org:', org, 'Suborg:', suborg, 'Subsuborg:', subsuborg, 'Address', address, 'Email', email
                        ###Assign values

                        authors_dict[author]={'org':org, 'suborg':suborg, 'subsuborg':subsuborg, 'address':address, 'email':email}


                        ###end of while for author address loop
            
                    
                ####Add abstracts to markups dict in loop (direct assignment did not work)

                Markups[abstract_reference]['Authors']={}  ##Put authors dict into markup
                markup['Authors']={} ##Include in loop (vestigal of early version)
                for metadata in markup.keys():
                    if metadata == 'Authors':
                        for author in authors_dict.keys():
                            Markups[abstract_reference]['Authors'][author]={}
                            for information in authors_dict[author].keys():
                                Markups[abstract_reference]['Authors'][author][information]=authors_dict[author][information]
                                #print abstract_reference, author, information, Markups[abstract_reference]['Authors'][author][information]
                ###end of author assignment    
                #print Markups[abstract_reference]['Authors'].keys()
                #print '\n',abstract_reference, Markups[abstract_reference]['Title']
                #print '\n',abstract_reference, Markups[abstract_reference]['Authors']
            ##try:
            ##    print Markups['kcl.ea.09.2016.SK.04']['Title']
            ##except:
            ##    pass
        ###end of authors infomrmation    
        else:
            pass ## placeholder for the rest of information.
        
        


    ###Open and assign workbook
    wb=Workbook()
    ws1=wb.active
    ws1.title="Papers"
    ws1=wb["Papers"]
    ws2=wb.create_sheet("AuthOrg")

    ws1['A1']='Accession Number'
    ws1['B1']='Title'
    ws1['C1']='Abstract'
    ws1['D1']='Language'
    ws1['E1']='Source'
    ws1['F1']='Date Value'

    ws2['A1']='Accession Number'
    ws2['B1']='Author'
    ws2['C1']='Country'
    ws2['D1']='Main Org'
    ws2['E1']='SubOrg'
    ws2['F1']='SubSubOrg'
    ws2['G1']='Email'
    ws2['H1']='Telephone'
    ws2['I1']='Address'

    excelfile='%sKCL_markups_%s_%s_%02d.xlsx' % ( Myfilepath, Region, Year, Months_no_dict[Month])
    row=1
    author_row=1    
    
    ###Loop over markups and write xlsx file
    abstract_list=Markups.keys()
    abstract_list.sort()
    

    for accession_no in abstract_list:
        print 'Adding ', accession_no, Markups[accession_no]['Title']

        ###print abstract , Markups[accession_no]['Title'], Markups[accession_no]['Language'], Markups[acession_no]['Date'] 

        row+=1
        cell='A%d' % row
        ws1[cell] = accession_no
        try:
            cell='B%d' % row
            ws1[cell] = Markups[accession_no]['Title']
        except:
            print 'Could not find Title for Accession number ', accession_no
            exit(2)
        try:    
            cell='D%d' % row
            ws1[cell] = Markups[accession_no]['Language']
        except:
            print '******Could not find Language for ', Markups[accession_no]['Title']
        try:    
            cell='E%d' % row
            ws1[cell] = Markups[accession_no]['Source']
        except:
            print '******Could not find Source for ', Markups[accession_no]['Title']
        try:    
            cell='F%d' % row
            ws1[cell] = Markups[accession_no]['Date']
        except:
            print '******Could not find Date for ', Markups[accession_no]['Title']
        if 'Authors' in Markups[accession_no].keys():

            for author in Markups[accession_no]['Authors'].keys():
                author_row+=1

                cell='A%d' % author_row
                ws2[cell] =  accession_no
                cell='B%d' % author_row
                ws2[cell] =  author
                cell='C%d' % author_row
                ws2[cell] =  Markups[accession_no]['Country']
                cell='D%d' % author_row
                ws2[cell] =  Markups[accession_no]['Authors'][author]['org']
                cell='E%d' % author_row
                ws2[cell] =  Markups[accession_no]['Authors'][author]['suborg']
                cell='F%d' % author_row
                ws2[cell] =  Markups[accession_no]['Authors'][author]['subsuborg']
                cell='GI%d' % author_row
                ws2[cell] =  Markups[accession_no]['Authors'][author]['email']
                cell='I%d' % author_row
                ws2[cell] =  Markups[accession_no]['Authors'][author]['address']
        
        else:
            author_row+=1
            ###print accession_no, author
            cell='A%d' % author_row
            ws2[cell] =  accession_no
            
            
            ###print author,  Markups[accession_no]['Authors'][author]['org'],  Markups[accession_no]['Authors'][author]['email'],  Markups[accession_no]['Authors'][author]['address']

    ###Write and save the file        
    wb.save(excelfile)
    print 'Wrote to', excelfile


if __name__=="__main__":
    main(sys.argv[1:])
