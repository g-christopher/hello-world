#!/usr/bin/enAv python
import docx  ##for microsoft word
import sys
import getopt
import os.path
import time
from openpyxl import Workbook  ##for microsoft excel
from strippers import strip_address ##strips institution and address
from strippers import strip_email ##strips {} email addresses
from strippers import strip_language ##strips () language
from strippers import strip_title
from strippers import aucleanup ##removes author's titles
from crossrefAPI import Getdoiplus ###parses
from crossrefAPI import GetFullAuthors ###parses full authors from doi

def entity_reader(Myfilepath, report_name, Month, Year, getdoi, debug):
    abstracts_path=Myfilepath
    
    #report_name='KSA'
    #Month='February'
    #Year='18'
    Months_date={'January':'31 Jan', 'February':'28 Feb', 'March': '31 March', 'April':'30 April', 'May':'31 May', 'June': '30 June', 'July':'31 July', 'August':'31 August', 'September':'30 Sept', 'October':'31 Oct', 'November':'30 Nov', 'December':'30 Dec'}
    Months_no_dict={'January':1, 'February':2, 'March':3, 'April':4, 'May':5, 'June':6, 'July':7, 'August':8, 'September':9, 'October':10, 'November':11, 'December':12}
    Months_no_short_dict={'Jan':1, 'Feb':2, 'Mar':3, 'Apr':4, 'May':5, 'Jun':6, 'Jul':7, 'Aug':8, 'Sep':9, 'Oct':10, 'Nov':11, 'Dec':12}
    #month_no=Months_no_short_dict[Month]
    for name in Months_no_dict.keys():
        if Month in name:
            month_no=Months_no_dict[name]
        

    ### Check for debug mode
    if debug==1:
        print '*********\n'
        print 'Debug mode'
        for i in xrange(0,5):
            time.sleep(1)
            sys.stdout.write('...\r')
            sys.stdout.flush()
        print '\n'

    write_excel=1
    ###Check if Endnote file is written
    EndNotefilepath='/Users/itadmin/Documents/IAEA/Dropbox_Mirror/EndNotedump/'
    EndNotereportname=report_name.replace(' ','_')
    EndNotefilename='%s%s_%s_%02d.txt' % ( EndNotefilepath, EndNotereportname, Year, month_no)
    if os.path.isfile(EndNotefilename) and debug==0:
        print 'Already written file ', EndNotefilename
        return

    ###Error File
    ### Write DOI files to import to endnote                                                                                                                                                      
    Errfilepath='/Users/itadmin/Documents/Agency'
    Errreportname=report_name.replace(' ','_')
    Errfilename='%s%s_%s_%02d.log' % ( Errfilepath, Errreportname, Year, month_no)
    err_file_= open(Errfilename,"w")
    print 'Writing Missed DOis to', Errfilename

    err_file_.write(Errreportname)
    


    try:
        doc = docx.Document(abstracts_path)
        print "Opening ", abstracts_path
        length =  len(doc.paragraphs)
    except IOError:
        print 'Cannot open: ', abstracts_path, '\n Check filepath and file exists'
    ###Define variables

    Country=''
    abstract_number=0
    Markups={} ##dict for all markups
    markup={} ##dummy dict for each markup
    abstract_reference=''
    authors_dict={} ##dummy dict for authors
    authorsinc=0  ##to include authors over multiple lines
    doiinc=0 ## For titles where the doi is put in

    Country_index='EN'
    Language='ENGLISH'
    Country='England'
    abstract_number=0
    ###Loop through paragraphs in document
    Title=''


    for para in doc.paragraphs:
        para.text = para.text.replace(u'\xa0',u' ')
        #print para.text
        if 'Title:' in para.text or 'Topic:' in para.text:
            ###check for Country
            try:
                Country
            except:
                print 'Country Section Heading Missing'
                exit(2)
            ###Get new abstract number
            abstract_number=abstract_number+1
            authorsinc=0
            FirstAuthorinc=0
            doiinc=0
            abstract_reference='ri.%s.%02d.20%s.%02d' % (report_name, month_no, Year, abstract_number)
            abstract_reference=abstract_reference.replace(" ","_")
            Markups[abstract_reference]={}
            ###print para.text.strip()
            ###Empty markup and authors dict
            markup.clear()
            authors_dict.clear()
            Title=para.text
            for line in para.text.splitlines():
                if 'Title:' in line or 'Topic:' in line:
                    Title=line

            Title=Title.replace('Title:',"",1)
            ###Strip unicode quotation marks from Title

            if u'\u201C' in Title or u'\u2018' in Title:
                Title=strip_title(Title)
                ###print 'New Title=', Title
            Title=Title.replace('Topic:',"",1)
            
            ###print Language, Title
            Language='ENGLISH'
            Lang_dict=['INDONESIAN', 'MANDARIN', 'ARABIC', 'FRENCH', 'TURKISH', 'FARSI', 'KOREAN', 'VIETNAMESE', 'JAPANESE', 'SPANISH']
            for language_ in Lang_dict:
                if language_ in Title:
                    print '*** detecting language',
                    Language=language_
                    print ' -- Language=', Language
                    Title=Title.replace('(%s)' % Language,"",1)
            

            Title=Title.strip()
            ###print Language, Title    
            markup['Title']=Title
            markup['Language']=Language
            markup['Country']=Country
            print abstract_reference, Title
            Markups[abstract_reference]['Title']=Title
            Markups[abstract_reference]['Language']=Language
            Markups[abstract_reference]['Country']=Country
            
            
        if 'Date:' in para.text and Title !='':
            Date=para.text
            if len(para.text.splitlines())>1:
                for line in para.text.splitlines():
                    if 'Date:' in line:
                        Date=line
            Date=Date.replace('Source Date:',"",1)
            Date=Date.replace('Date:',"",1)
            Date=Date.replace('Accessed',"",1)
            Date=Date.strip()
            markup['Date']=Date
            ###print "Date=",  Date, '\n'
            Markups[abstract_reference]['Date']=Date
            
        if 'Source:' in para.text or 'source:' in para.text or 'Original source:' in para.text or 'Journal:'in para.text:    

            Source=para.text
            if len(para.text.splitlines())>1:
                for line in para.text.splitlines():
                    if 'Source:' in para.text or 'source:' in para.text or 'Original source:' in para.text or 'Journal:'in para.text:    
                        Date=line
            Source=Source.replace('Original Source:',"",1)
            Source=Source.replace('Original source:',"",1)
            Source=Source.replace('Source:',"",1)
            Source=Source.replace('Journal:',"",1)
            Source=Source.strip()
            Source=Source.replace('(%s)' % Language, "",1)
            markup['Source']=Source
            Markups[abstract_reference]['Source']=Source
            #print Source.split(',')
            markup['Volume']=''
            markup['Issue']=''

            ###replacement words
            voldict=['Vol.','volume','vol.','Volume']
            issdict=['Issue','Iss.','issue','iss.']
            if len(Source.split(','))>1:
                markup['Journal']=Source.split(',')[0]
                for volname in voldict:
                    if volname in markup['Journal']:
                        markup['Journal']= markup['Journal'].rsplit(volname,1)[0]
                        ###print 'Journal is', markup['Journal']
            for thingy in Source.split(','):
                for volname in voldict:
                    if volname in thingy:
                        try:
                            markup['Volume']= thingy.strip().rsplit(volname,1).pop()
                            ###print 'Volume=', thingy.strip().rsplit(volname,1).pop()
                        except:
                            print 'Cannot resolve volume for', Title
                        
                    #print markup['Volume']
                for issname in issdict:    
                    if issname in thingy:
                        try:
                            markup['Issue']= thingy.strip().rsplit(issname,1).pop()
                            ###print 'Issue=', thingy.strip().rsplit(issname,1).pop()
                        except:
                            print 'Cannot resolve issue for', Title
                        
                    #print markup['Issue']
                
            ##for line in Source.split(','):
            ##    if 'vol' in line.lower():
            ##        markup['Volume']=line.split().strip()
            ###print 'Source is', Source

            ###add markup metadata
            Markups[abstract_reference]={}
            for metadata in markup.keys():
                try:
                    Markups[abstract_reference][metadata]=markup[metadata]
                    #print 'Adding', metadata
                except:
                    Markups[abstract_reference][metadata]=''
                    print 'No entry for ',metadata
        if ('doi\.org' in para.text.lower() or 'DOI:' in para.text) and doiinc==0:
            doiinc=1
            doi=''
            if 'doi\.org' in para.text.lower():
                doi=para.text.split('doi\.org/').pop()
                if ' ' in doi:
                    doi=doi.split(' ')[0].strip()
            elif 'DOI:' in para.text:
                doi=para.text.split('DOI:').pop().strip()
            markup['DOI']=doi
            Markups[abstract_reference]['DOI']=markup['DOI']
            print '*** Found doi in text: ', Markups[abstract_reference]['DOI']
            
        if 'Authors:' in para.text or 'Author:' in para.text or 'affiliation:' in para.text.lower() or 'affiliations:' in para.text.lower() or 'Researchers:' in para.text or 'authors were:' in para.text or 'author was:' in para.text or authorsinc==1:
            ###control for authors over multiple lines

            if authorsinc>0 and not para.text:
                authorsinc=0 ### terminate if empty line

            elif authorsinc>0 and 'Abstract:' in para.text:
                authorsinc=0 ### terminate if Abstract is on the nextline

            else:
                if authorsinc==0:
                    authorsinc=1+authorsinc ### Increment authorsinc so will contine to get data from next lines if zero.
                    FirstAuthorinc=0 ### Set to zero to get first author
                    #print 'resetting first author inc'
                ###Define dummy variables for author dicts
                authors_dummy_list=[]
                multi_addr_list=[]
                address=''
                email=''
                org=''
                suborg=''
                subsuborg=''
                authors_dummy=''
                print 'In Authors'

                ###Strip intial text
                aulist=['Author/affiliation:', 'Authors:', 'Author:', 'Authors/Affiliation:', 'Authors/affiliation:', 'Author/Affiliation:', 'Authors/affiliations:', 'Researchers:', 'Authors/Affiliations:','authors were:' , 'author was:']
                for auname in aulist:
                    if auname in para.text:
                        authors_dummy=para.text.split(auname).pop().strip()

                        
                ###Look for author addresses and loop over authors
                

                ###Find first author using ICSA's multiple methods of splitting the author list
                if FirstAuthorinc==0:
                    FirstAuthorinc=1
                    ###some seperators are ; some are ,

                    fadict={'from':authors_dummy.find('from'),'(':authors_dummy.find('('),',':authors_dummy.find(','), ';':authors_dummy.find(';'),'based':authors_dummy.find('based'), ' and ':authors_dummy.find(' and '), 'Associate':authors_dummy.find('Associate'), 'Assistant':authors_dummy.find('Assistant'), 'is based':authors_dummy.find('is based') , '{': authors_dummy.find('{'), '|': authors_dummy.find('|'), ']':authors_dummy.find(']')}
                    ###decide to use comma as delimiter
                    if ',' in authors_dummy and ';' in authors_dummy:
                        #print 'distance is', abs(authors_dummy.find(';')-authors_dummy.find(','))
                        if abs(authors_dummy.find(';')-authors_dummy.find(','))<15:
                            dummy=fadict.pop(',', None)
                    if ',' in authors_dummy and '{' in authors_dummy:

                        if abs(authors_dummy.find('{')-authors_dummy.find(','))<15:
                            dummy=fadict.pop(',', None)
                    if ',' in authors_dummy and '(' in authors_dummy:

                        if abs(authors_dummy.find('(')-authors_dummy.find(','))<15:
                            dummy=fadict.pop(',', None)        
                    if '|' in authors_dummy or ']' in authors_dummy: ###Endnote 8 delimiters
                            dummy=fadict.pop(',', None)        
                            

                    if max(n for n in fadict.values()) <0:
                        ##print authors_dummy
                        ##print aucleanup(authors_dummy)
                        markup['FirstAuthor'] = aucleanup(authors_dummy)
                        ###control for n/a authors
                        if aucleanup(authors_dummy).lower()=='n/a':
                            markup['FirstAuthor']=''
                        
                        #print 'no from, bracket, comma, based or and'
                    else:    
                        lowest=min(n for n in fadict.values() if n>0)
                        for symbol in fadict.keys():
                            if lowest==fadict[symbol]:
                                #print 'found a symbol at', lowest, authors_dummy[lowest], authors_dummy[lowest+1], authors_dummy[lowest+2]
                                #print authors_dummy.split(symbol)
                                #print authors_dummy.split(symbol)[0]
                                #print authors_dummy.split(symbol)[0].strip()
                                ##print authors_dummy.split(symbol)[0].strip()
                                ##print aucleanup(authors_dummy.split(symbol)[0].strip())
                                markup['FirstAuthor'] = aucleanup(authors_dummy.split(symbol)[0].strip())
                    if ',' in markup['FirstAuthor']:
                        markup['FirstAuthor']=markup['FirstAuthor'].split(',')[0].strip()
                    print '%-150s First Author= %s' % (Title, markup['FirstAuthor'])    
                    Markups[abstract_reference]['FirstAuthor']=markup['FirstAuthor']
                    #print authors_dummy.split(authors_dummy[lowest])[0].strip()


                    #Get Authors from crossref if English Language
                    if getdoi==1 and doiinc==0:
                        print 'Getting DOI'
                        try:
                            if markup['FirstAuthor']!='' and markup['Language']=='ENGLISH':
                                doi=Getdoiplus(markup)
                            
                            else:
                                doi=''
                        except:
                            doi=''
                            print '**** Failed to get Author from CrossRef'
                        doiinc=1
                        #except:
                        #    doi=''
                        #    print '**** Could not use CrossRef'
                        markup['DOI']=doi
                        Markups[abstract_reference]['DOI']=markup['DOI']
                        if doi=='':
                            try:
                                errstr=Markups[abstract_reference]['Title']+'\n'
                                err_file_.write(errstr)
                            except:
                                print '****Could not Write: ', errstr
                        ###authorsdict=GetFullAuthors(doi)
                    ### End of First Author inc
                    else:
                        print '\n**********\n'
                foobar=0
                while '(' in authors_dummy and foobar==1:
                    address= strip_address(authors_dummy)
                    authors_dummy=authors_dummy.replace('(%s)' % address,"",1)
                    authors_dummy_list=authors_dummy.split(',')
                    ###print 'address=',address, 'authors=',authors_dummy
                    ###Loop over authors by address
                    for author_dummy in authors_dummy_list:
                        email=''
                        author=author_dummy
                        
                        ###Find emails
                        if '{' in author_dummy:
                            #print 'finding email', author_dummy
                            email=strip_email(author_dummy)
                            author=author.replace('{%s}' % email,"",1)
                            email=email.strip('email:')
                            email=email.strip('{')
                            email=email.strip('}')
                            email=email.strip()
                            ###end of emails    

                        ###Multiple addresses
                        if '/' in address:
                            ###print 'multi addr', author
                            multi_addr_list=address.split('/')
                            ###print multi_addr_list
                        ###Asign orgs, suborgs, subsuborgs
                        author=author.strip()
                        orgs=address.split(',')
                        org=orgs[0]
                        org=org.strip()
                        
                        if len(orgs)==3:
                            org=orgs[2]
                            suborg=orgs[1]
                            subsuborg=orgs[0]
                        elif len(orgs)==2:
                            org=orgs[1]
                            suborg=orgs[0]
                            subsuborg=''
                        elif len(orgs)==1:
                            org=orgs[0]
                            suborg=''
                            subsuborg=''
                        else:
                            org=''
                            suborg=''
                            subsuborg=''
                            
                        org=org.strip()
                        suborg=suborg.strip()
                        subsuborg=subsuborg.strip()
                        ###print '\nauthor', author, 'org', org, 'suborg', suborg, 'subsuborg', subsuborg, 'address', address, 'email', email
                        ###Assign values

                        authors_dict[author]={'org':org, 'suborg':suborg, 'subsuborg':subsuborg, 'address':address, 'email':email}


                        ###end of while for author address loop
            
                    
                ####Add abstracts to markups dict in loop (direct assignment did not work)
                
                #Markups[abstract_reference]['Authors']={}  ##Put authors dict into markup
                #markup['Authors']={} ##Include in loop (vestigal of early version)
                #for metadata in markup.keys():
                #    if metadata == 'Authors':
                #        
                #        for author in authors_dict.keys():
                #            Markups[abstract_reference]['Authors'][author]={}
                #            for information in authors_dict[author].keys():
                #                Markups[abstract_reference]['Authors'][author][information]=authors_dict[author][information]
                ###end of author assignment    
            
            ###print '\n',abstract_reference, Markups[abstract_reference]['Title']
            ###print '\n',abstract_reference, Markups[abstract_reference]['Authors']
            ##try:
            ##    print Markups['kcl.ea.09.2016.SK.04']['Title']
            ##except:
            ##    pass
        ###end of authors infomrmation    
        else:
            pass ## placeholder for the rest of information.
        
        

    if write_excel==1 and debug==0:    
        ###Open and assign workbook
        wb=Workbook()
        ws1=wb.active
        ws1.title="Papers"
        ws1=wb["Papers"]
        ws2=wb.create_sheet("AuthOrg")

        ws1['A1']='Accession Number'
        ws1['B1']='Title'
        ws1['C1']='Abstract'
        ws1['D1']='Language'
        ws1['E1']='Source'
        ws1['F1']='Journal'
        ws1['G1']='Volume'
        ws1['H1']='Issue'
        ws1['I1']='Date Value'
        ws1['J1']='First Author'
        ws1['K1']='DOI'
        
        ws2['A1']='Accession Number'
        ws2['B1']='Author'
        ws2['C1']='Country'
        ws2['D1']='Main Org'
        ws2['E1']='SubOrg'
        ws2['F1']='SubSubOrg'
        ws2['G1']='Email'
        ws2['H1']='Telephone'
        ws2['I1']='Address'
    
        Excelfilepath='/Users/itadmin/Documents/Agency'
        Excelreportname=report_name.replace(' ','_')
        excelfile='%s%s_%s_%02d.xlsx' % ( Excelfilepath, Excelreportname, Year, month_no)
        row=1
        author_row=1    
    
        ###Loop over markups and write xlsx file
        abstract_list=Markups.keys()
        abstract_list.sort()

        for accession_no in abstract_list:
            ###print 'Adding ', accession_no, Markups[accession_no]['Title']

            ###print abstract , Markups[accession_no]['Title'], Markups[accession_no]['Language'], Markups[acession_no]['Date'] 

            row+=1
            cell='A%d' % row
            ws1[cell] = accession_no
            try:
                cell='B%d' % row
                ws1[cell] = Markups[accession_no]['Title']
            except:
                print 'Could not find Title for Accession number ', accession_no, '*****'
                exit(2)
            try:    
                cell='D%d' % row
                ws1[cell] = Markups[accession_no]['Language']
            except:
                print '******Could not find Language for ', accession_no, Markups[accession_no]['Title'], '*****'
            try:    
                cell='E%d' % row
                ws1[cell] = Markups[accession_no]['Source']
            except:
                pass
                ###print '******Could not find Source for ', accession_no, Markups[accession_no]['Title'], '*****'
            try:    
                cell='F%d' % row
                ws1[cell] = Markups[accession_no]['Journal']
            except:
                pass
            try:    
                cell='G%d' % row
                ws1[cell] = Markups[accession_no]['Volume']
            except:
                pass
            try:    
                cell='H%d' % row
                ws1[cell] = Markups[accession_no]['Issue']
            except:
                pass
            try:    
                cell='I%d' % row
                ws1[cell] = Markups[accession_no]['Date']    
            except:
                pass
                ###print '******Could not find Date for ', accession_no, Markups[accession_no]['Title'], '*****'
            try:    
                cell='J%d' % row
                ws1[cell] = Markups[accession_no]['FirstAuthor']
            except:
                print '******Could not find First Author for ', accession_no, Markups[accession_no]['Title'], '*****'
            try:
                cell='K%d' % row
                print Markups[accession_no]['DOI']
                ws1[cell] = Markups[accession_no]['DOI']
            except:
                print '******Could not find DOI for ', accession_no, Markups[accession_no]['Title'], '*****'
            if 'Authors' in Markups[accession_no].keys():
                for author in Markups[accession_no]['Authors'].keys():
                    author_row+=1
                    ###print accession_no, author
                    cell='A%d' % author_row
                    ws2[cell] =  accession_no
                    cell='B%d' % author_row
                    ws2[cell] =  author
                    cell='C%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Country']
                    cell='D%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Authors'][author]['org']
                    cell='E%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Authors'][author]['suborg']
                    cell='F%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Authors'][author]['subsuborg']
                    cell='GI%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Authors'][author]['email']
                    cell='I%d' % author_row
                    ws2[cell] =  Markups[accession_no]['Authors'][author]['address']
        
            else:
                author_row+=1
                ###print accession_no, author
                cell='A%d' % author_row
                ws2[cell] =  accession_no
            
            
                ###print author,  Markups[accession_no]['Authors'][author]['org'],  Markups[accession_no]['Authors'][author]['email'],  Markups[accession_no]['Authors'][author]['address']

        ###Write and save the file        
        wb.save(excelfile)
        print 'Wrote to', excelfile

    if getdoi==1 and debug==0:    
        ### Write DOI files to import to endnote
        EndNotefilepath='/Users/itadmin/Documents/Agency'
        EndNotereportname=report_name.replace(' ','_')
        EndNotefilename='%s%s_%s_%02d.txt' % ( EndNotefilepath, EndNotereportname, Year, month_no)
        doi_file_= open(EndNotefilename,"w")
        abstract_list=Markups.keys()
        abstract_list.sort()
        
        print 'Writing DOIs to', EndNotefilename
        for accession_no in abstract_list:
            if 'DOI' in Markups[accession_no].keys():
                if Markups[accession_no]['DOI']!='':
                    doi_line_str='DOI: %s | %s\n' %  (Markups[accession_no]['DOI'], Markups[accession_no]['Title'])
                    if u'\u2019' in doi_line_str:
                        doi_line_str=doi_line_str.replace(u'\u2019','\'')
                        print 'replacing \''
                    if u'\u2013' in doi_line_str:
                        doi_line_str=doi_line_str.replace(u'\u2013','\'')
                        print 'replacing \''
    
                    doi_file_.write(doi_line_str)    
                    
                    #try:
                    #
                    #except:
                    #    print '*** cannot write doi'
                    #    #print 'Writing', doi_line_str
        doi_file_.close()
        
    err_file_.close()
    return
